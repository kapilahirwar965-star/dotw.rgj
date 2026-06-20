# -- coding: utf-8 --
import sys
import io
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import streamlit as st
import os
import io
import datetime
from PIL import Image
import pypdf
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
from openai import OpenAI

# Try importing the new Google GenAI SDK, fallback to the legacy google-generativeai
try:
    from google import genai
    from google.genai import types
    HAS_NEW_SDK = True
except ImportError:
    import google.generativeai as genai_legacy
    HAS_NEW_SDK = False

# ---------------------------------------------------------
# Intelligent Department Router Configuration
# ---------------------------------------------------------
ROUTING_MAP = {
    "आंकड़े / सांख्यिकी / जनगणना / विकासखंडवार संकलित जानकारी / अनुसूचित क्षेत्र घोषणा डेटा": "जिला सांख्यिकी अधिकारी, राजगढ़",
    "निर्माण / भवन / मरम्मत / अधोसंरचना / तकनीकी स्वीकृति": "कार्यपालन यंत्री, ग्रामीण यांत्रिकी सेवा (RES), राजगढ़",
    "छात्रावास / प्रवेश / छात्रवृत्ति / जनजातीय कल्याण / आश्रम": "सहायक आयुक्त, जनजातीय कार्य विभाग, राजगढ़",
    "कोर्ट केस / लीगल / प्रभारी अधिकारी नियुक्ति / जवाब दावा": "शासकीय अधिवक्ता / शाखा प्रभारी (विधि), राजगढ़"
}

# म.प्र. शासन कार्यप्रणाली के अनुसार डायनेमिक एड्रेसिंग गाइडलाइन
DYNAMIC_ROUTING_INSTRUCTION = """
    आप मध्य प्रदेश शासन के प्रशासनिक नियमों (Administrative Rules of Business) के विशेषज्ञ हैं।
    प्रारूप (Draft) तैयार करते समय 'प्रति' (To Address) का निर्धारण नीचे दिए गए नियमों के अनुसार पूरी तरह डायनेमिक होना चाहिए:
    1. पत्र के मुख्य विषय, संदर्भ और आंतरिक पाठ (Internal Text) को गहराई से पढ़ें।
    2. पत्र जिस विशिष्ट अधिकारी या पद (Designation) के लिए भेजा जाना है, उसका सटीक पदनाम ही 'प्रति' ब्लॉक में अंकित करें।
    3. यदि पत्र में स्पष्ट रूप से किसी ब्लॉक अधिकारी, छात्रावास अधीक्षक, लोक निर्माण विभाग (PWD) के यंत्री, या राज्य स्तर के आयुक्त का संदर्भ है, तो सीधे जिला प्रमुख का नाम लिखने के बजाय उसी सटीक पद का संबोधन लिखें।
    4. शासकीय मर्यादा और पदानुक्रम (Hierarchy) का पूर्ण पालन होना चाहिए।
"""

# ---------------------------------------------------------
# Page Configuration & Styling (Premium MP Govt Theme)
# ---------------------------------------------------------
st.set_page_config(
    page_title="मध्य प्रदेश शासन - राजगढ़ शासकीय प्रारूपक",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for theme & aesthetics
st.markdown("""
<style>
    /* Primary brand colors: Deep Maroon, Saffron, Gold, Light Sand */
    :root {
        --primary-maroon: #7A1C1C;
        --secondary-saffron: #E25822;
        --accent-gold: #D4AF37;
        --bg-light: #F9F6F0;
    }
    
    /* Global styles */
    .stApp {
        background-color: #F9F6F0;
        font-family: 'Inter', 'Noto Serif Devanagari', 'Segoe UI', serif;
    }
    
    /* Header styling */
    .header-container {
        background: linear-gradient(135deg, #7A1C1C 0%, #B83A3A 100%);
        color: white;
        padding: 24px;
        border-radius: 12px;
        margin-bottom: 24px;
        border-bottom: 4px solid #D4AF37;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        text-align: center;
    }
    .header-title {
        font-size: 2.2rem;
        font-weight: bold;
        margin: 0;
        color: #FFFFFF;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.5);
    }
    .header-subtitle {
        font-size: 1.1rem;
        margin-top: 8px;
        color: #FFE5B4;
        letter-spacing: 1px;
    }
    
    /* Card/Section styling */
    .css-card {
        background-color: white;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        margin-bottom: 20px;
        border-left: 5px solid #7A1C1C;
    }
    
    /* Buttons */
    .stButton>button {
        background: linear-gradient(135deg, #7A1C1C 0%, #9E2A2A 100%);
        color: white !important;
        border: none;
        padding: 10px 24px;
        border-radius: 6px;
        font-weight: bold;
        transition: all 0.3s ease;
        box-shadow: 0 2px 5px rgba(0,0,0,0.15);
    }
    .stButton>button:hover {
        background: linear-gradient(135deg, #E25822 0%, #FF7A45 100%);
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.25);
    }
    
    /* Labels and Headings */
    h1, h2, h3 {
        color: #7A1C1C;
    }
    
    /* Info alerts */
    .stAlert {
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# Sidebar Setup (Customized for Rajgarh Collectorship)
# ---------------------------------------------------------
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; margin-bottom: 20px;">
        <span style="font-size: 4rem;">🏛️</span>
        <h3 style="margin-top: 10px; color: #7A1C1C; font-weight: bold;">राजगढ़ शासकीय प्रारूपक</h3>
        <p style="font-size: 0.85rem; color: #666;">जिला राजगढ़ (ब्यावरा) म.प्र.</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.header("⚙️ विन्यास (Configuration)")
    
    gemini_key = st.text_input(
        "गूगल जेमिनी API कुंजी (Gemini API Key)",
        type="password",
        placeholder="AIzaSy-Key",
        key="gemini_key_input",
        help="जेमिनी का उपयोग करने के लिए अपनी API कुंजी दर्ज करें।"
    )

    claude_key = st.text_input(
        "एंथ्रोपिक क्लॉड API कुंजी (Claude API Key)",
        type="password",
        placeholder="sk-ant-...",
        key="claude_key_input",
        help="क्लॉड का उपयोग करने के लिए अपनी API कुंजी दर्ज करें।"
    )

    openai_key = st.text_input(
        "OpenAI API कुंजी (OpenAI API Key)",
        type="password",
        placeholder="sk-proj-...",
        key="openai_key_input",
        help="OpenAI का उपयोग करने के लिए अपनी API कुंजी दर्ज करें।"
    )

    provider = st.selectbox(
        "मॉडल प्रदाता (Model Provider)",
        options=[
            "Google Gemini",
            "Anthropic Claude",
            "OpenAI"
        ],
        index=0,
        key="model_provider_selection",
        help="पत्र जनरेट करने के लिए सेवा प्रदाता का चयन करें।"
    )

    if provider == "Google Gemini":
        selected_model = st.selectbox(
            "जेमिनी मॉडल (Gemini Model)",
            options=[
                "gemini-2.5-pro",
                "gemini-2.5-flash",
                "gemini-2.5-flash-lite",
                "gemini-1.5-pro"
            ],
            index=0,
            key="gemini_model_selection",
            help="शासकीय प्रारूप तैयार करने के लिए जेमिनी मॉडल का चयन करें।"
        )
    elif provider == "Anthropic Claude":
        selected_model = st.selectbox(
            "क्लॉड मॉडल (Claude Model)",
            options=[
                "claude-3-5-sonnet"
            ],
            index=0,
            key="claude_model_selection",
            help="शासकीय प्रारूप तैयार करने के लिए क्लॉड मॉडल का चयन करें।"
        )
    else:
        selected_model = st.selectbox(
            "OpenAI मॉडल (OpenAI Model)",
            options=[
                "gpt-4o"
            ],
            index=0,
            key="openai_model_selection",
            help="शासकीय प्रारूप तैयार करने के लिए OpenAI मॉडल का चयन करें।"
        )
    
    st.divider()
    
    st.header("📝 शासकीय पत्र विनिर्देश")
    
    workflow_type = st.selectbox(
        "पत्राचार प्रवाह प्रकार (Workflow Type)",
        options=[
            "वरिष्ठ कार्यालय को अंतिम उत्तर/प्रतिवेदन (Final Outward Reply)",
            "जवाब देही / अंतर्विभागीय जानकारी मांग पत्र (Inter-Department Inquiry)",
            "अधीनस्थ को निर्देश/पत्र (District/Subordinate Order)",
            "कार्यालयीन टिप्पणी/नोटशीट (Office Notesheet)",
            "सामान्य शासकीय पत्र (Official Letter)",
            "कार्यालय आदेश (Office Order)"
        ],
        index=0
    )
    
    dept_header = "कार्यालय कलेक्टेरेट (जनजातीय कार्य तथा अनुसूचित जाति कल्याण विभाग)"
    district_info = "जिला राजगढ़ (ब्यावरा) म.प्र."
    email_id = "Dotw.rjg@mp.gov.in"
    
    branch_name = st.selectbox(
        "शाखा का नाम (Select Department Branch)",
        options=[
            "योजना",
            "छात्रावास",
            "स्थापना",
            "निर्माण",
            "PMAGY (PM-AJAY)",
            "भवन किराया निर्धारण/भुगतान",
            "कोर्ट केस",
            "RTI",
            "विविध"
        ],
        key="selected_branch"
    )
    outward_no = st.text_input("जावक क्रमांक (Outward Number)", value=f"क्रमांक /       / {branch_name} / 2026-27")
    
    # यहाँ सुधार किया गया है - वेरियबल को उपयोग से पहले बनाया गया है
    current_date_hi = datetime.datetime.now().strftime("%d/%m/%Y")
    letter_date = st.text_input("दिनांक (Date)", value=current_date_hi)
    
    signatory = st.selectbox(
        "हस्ताक्षरकर्ता अधिकारी (Signatory Officer)",
        options=[
            "कलेक्टर (Collector)",
            "जिला संयोजक (District Organizer)"
        ],
        index=0
    )
    
    footer_path = st.text_input("फाइल-पथ फुटर (File-Path Footer)", value="C:\\Users\\Desktop\\Draft.docx")

# ---------------------------------------------------------
# Main Page Header
# ---------------------------------------------------------
st.markdown("""
<div class="header-container">
    <div class="header-title">कार्यालय कलेक्ट्रेट (जनजातीय कार्य तथा अनुसूचित जाति कल्याण विभाग)</div>
    <div class="header-subtitle">शासकीय कार्यप्रणाली आलेखन सहायक - जिला राजगढ़ (ब्यावरा) मध्य प्रदेश</div>
</div>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------
def extract_text_from_pdf(pdf_file):
    """Extracts all text contents from an uploaded PDF file."""
    try:
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"पीडीएफ से टेक्स्ट निकालते समय त्रुटि: {str(e)}")
        return ""

def get_system_instruction(workflow, dept, district, email, branch, out_no, date, sign_off, f_path):
    """Builds a strict, detailed system instruction for the Gemini API with Rajgarh guidelines."""
    
    if "कलेक्टर" in sign_off:
        primary_sig = "कलेक्टर\nजिला राजगढ़ (म.प्र.)"
        subordinate_sig = "जिला संयोजक\nअनुसूचित जाति तथा जनजातीय कल्याण विभाग\nजिला राजगढ़ (म.प्र.)"
    else:
        primary_sig = "जिला संयोजक\nअनुसूचित जाति तथा जनजातीय कल्याण विभाग\nजिला राजगढ़ (म.प्र.)"
        subordinate_sig = "जिला संयोजक\nअनुसूचित जाति तथा जनजातीय कल्याण विभाग\nजिला राजगढ़ (म.प्र.)"

    branch_contexts = {
        "योजना": "योजना शाखा मुख्य रूप से जनजातीय कल्याण योजनाओं, सांख्यिकी, लक्ष्यों, आवंटन और प्रशासनिक रिपोर्टिंग पर केंद्रित है.",
        "छात्रावास": "छात्रावास शाखा मुख्य रूप से जनजातीय छात्रावास प्रबंधन, छात्र प्रवेश, बुनियादी सुविधाओं, छात्रावास अधीक्षक निगरानी और निरीक्षण पर केंद्रित है.",
        "स्थापना": "स्थापना शाखा विभाग के कर्मियों के सेवा अभिलेखों, पदोन्नति, अवकाश, स्थापना आदेश और स्थानांतरण मामलों पर केंद्रित है.",
        "निर्माण": "निर्माण शाखा इंजीनियरिंग कार्यों, बुनियादी ढांचे के विकास, स्वीकृतियों, भौतिक/वित्तीय प्रगति की निगरानी और पूर्णता प्रमाण पत्र (CC) के अनुमोदन पर केंद्रित है.",
        "PMAGY (PM-AJAY)": "PMAGY (PM-AJAY) शाखा प्रधान मंत्री अनुसूचित जाति अभ्युदय योजना, ग्राम विकास योजनाओं, व्यक्तिगत लाभों और अभिसरण (convergence) पर केंद्रित है.",
        "भवन किराया निर्धारण/भुगतान": "भवन किराया निर्धारण/भुगतान शाखा कार्यालय/छात्रावास के भवन किराए के मूल्यांकन, लोक निर्माण विभाग (PWD) के मूल्यांकन प्रमाण पत्रों, लीज की शर्तों और वित्तीय प्रसंस्करण पर केंद्रित है.",
        "कोर्ट केस": "कोर्ट केस शाखा कानूनी मामलों, प्रभारी अधिकारी (OIC) की नियुक्तियों के प्रारूपण, पैरावार टिप्पणियों (parawise comments), रिट याचिकाओं (WPs), अवमानना नोटिसों और महाधिवक्ता (AG) कार्यालय के साथ समन्वय पर केंद्रित है.",
        "RTI": "RTI शाखा सूचना का अधिकार आवेदनों, प्रथम अपीलों, समय-सीमा में निपटान और सूचना का अधिकार अधिनियम 2005 की धाराओं पर केंद्रित है.",
        "विविध": "विविध शाखा सामान्य अंतर-विभागीय परिपत्रों, जन शिकायतों और विविध प्रशासनिक कार्यों का प्रबंधन करती है."
    }
    branch_desc = branch_contexts.get(branch, "विविध प्रशासनिक कार्यों पर केंद्रित है.")

    instruction = f"""
    आप मध्य प्रदेश शासन के "कार्यालय कलेक्टर (जनजातीय कार्य तथा अनुसूचित जाति कल्याण विभाग) जिला राजगढ़ (ब्यावरा) म.प्र." में पदस्थ एक अत्यंत अनुभवी और कुशल सहायक ग्रेड-3 (Assistant Grade-3) कर्मचारी हैं. आपको शासकीय कार्यप्रणाली और सचिवालयीय आलेखन (Government Drafting) का गहन ज्ञान है.

    आपका कार्य दिए गए संदर्भों और निर्देशों के आधार पर विशुद्ध रूप से शासकीय प्रारूप (Draft) तैयार करना है.

    आपको निम्नलिखित नियमों का कड़ाई से पालन करना होगा:
    
    1. भाषा और शब्दावली (राजभाषा हिन्दी):
       - आपकी भाषा अत्यंत औपचारिक, सुस्पष्ट और शासकीय शब्दावली से युक्त होनी चाहिए.
       - आपको अनिवार्य रूप से निम्नलिखित शासकीय मुहावरों और वाक्यांशों का उपयोग करना है:
         * "उपरोक्त विषयांतर्गत संदर्भित पत्र का अवलोकन करने का कष्ट करें..."
         * "निर्देशानुसार लेख है कि..."
         * "जानकारी निरंक है." (यदि कोई जानकारी खाली या शून्य है)
         * "आगामी आवश्यक कार्यवाही हेतु सादर प्रेषित."
       - किसी भी अनौपचारिक, बोलचाल की हिन्दी या सीधे अंग्रेजी शब्दों का उपयोग न करें.

    2. विशिष्ट लेटरहेड संरचना (शीर्ष प्रारूप):
       सभी पत्रों/ज्ञापनों/प्रतिवेदनों के शीर्ष पर निम्न संरचना का हुबहू उपयोग करें:
       
       कार्यालय कलेक्टर (जनजातीय कार्य तथा अनुसूचित जाति कल्याण विभाग)
       जिला राजगढ़ (ब्यावरा) म.प्र.
       E-Mail ID: {email}
       __________
       {out_no}                                 राजगढ़, दिनांक...................

    3. पत्राचार प्रवाह, नियम एवं शाखा विशिष्ट संदर्भ (Workflow, Routing & Branch Context):
       - पत्र का प्रकार: "{workflow}"
       - वर्तमान संबंधित शाखा: "{branch}"
       - शाखा का कार्यक्षेत्र विवरण: {branch_desc}
       
       नियम: इस शाखा के विशिष्ट कार्यक्षेत्र, उद्देश्यों और शब्दावली का आलेख प्रारूपण (Drafting) और कार्यालयीन टिप्पणी (Notesheet / नस्ती) लिखते समय कड़ाई से पालन करें. इसका टोन मध्य प्रदेश Civil Service नियमों और शासकीय प्रारूपण मानकों के सर्वथा अनुकूल होना चाहिए.
    """
    
    if "अंतिम उत्तर/प्रतिवेदन" in workflow:
        instruction += """
       - कार्य: यह वरिष्ठ कार्यालय (जैसे भोपाल संचालनालय) को भेजा जाने वाला अंतिम उत्तर है.
       - प्रति (To): पत्र का 'प्रति' (To) हमेशा आवक पत्र जारी करने वाले वरिष्ठ कार्यालय के नाम होना चाहिए (जैसे: 'प्रति, कार्यालय आयुक्त, जनजातीय कार्य, मध्य प्रदेश, भोपाल')।
       - विषय (Subject): पत्र का विषय स्वचालित रूप से '...की जानकारी उपलब्ध कराने बाबत्' से बदलकर '...की जानकारी/प्रतिवेदन प्रेषण बाबत्' होना चाहिए।
       - आपको पूर्व में प्राप्त सभी संबंधित पत्रों/आदेशों का संदर्भ (संदर्भ संख्या एवं दिनांक) विषय के ठीक नीचे 'संदर्भ:-' खंड में स्पष्ट रूप से देना होगा.
       - अधीनस्थ या अन्य कार्यालयों से प्राप्त जानकारी को एकत्रित कर एक संयुक्त प्रतिवेदन या अंतिम प्रतिउत्तर तैयार करें.
       - पत्र की मुख्य विषयवस्तु में "संदर्भित पत्र के अनुक्रम में जानकारी संलग्न प्रेषित है..." या "उपरोक्त विषयांतर्गत निर्देशानुसार लेख है कि..." का प्रयोग करें.
         """
    elif "intent-based" in workflow or "मतभेद" in workflow or "मतभेद" in workflow or "अंतर्विभागीय जानकारी मांग" in workflow:
        instruction += """
       - कार्य: अन्य स्थानीय शासकीय विभागों से जानकारी प्राप्त करने हेतु पत्र.
       - संबोधन में 'प्रति,' के अंतर्गत बुद्धिमान विभाग राउटर (Intelligent Department Router) द्वारा निर्धारित विभाग/अधिकारी का पदनाम और पता लिखें.
       - पत्र में समय-सीमा (Time-limit) का उल्लेख करें और त्वरित जानकारी भेजने का अनुरोध करें.
         """
    elif "अधीनस्थ को निर्देश" in workflow:
        instruction += """
       - कार्य: विकासखंड स्रोतों, छात्रावास अधीक्षकों, या क्षेत्र संयोजकों को निर्देश देना.
       - मुख्य वाक्य संरचना में "निर्देशित किया जाता है कि तत्काल..." या "समय-सीमा में जानकारी उपलब्ध कराना सुनिश्चित करें" का प्रयोग करें.
         """
    elif "नोटशीट" in workflow:
        instruction += """
       - कार्य: कार्यालयीन टिप्पणी (Notesheet) का प्रारूप.
       - इसमें शीर्ष पर बीच में "|| कार्यालयीन टिप्पणी ||" लिखें.
       - बाईं ओर 'विषय' लिखें. इसके बाद 'टीप क्रमांक 1' से शुरू करते हुए मामले का संक्षिप्त विवरण, प्राप्त पत्र का विवरण और सक्षम अधिकारी (जैसे कलेक्टर महोदय) से आदेश प्राप्त करने हेतु टिप्पणी तैयार करें.
       - अंत में दाईं ओर सहायक ग्रेड-3 के हस्ताक्षर का स्थान तथा उसके नीचे क्रमिक अधिकारियों (जिला संयोजक -> अपर कलेक्टर -> कलेक्टर) के हस्ताक्षर के लिए पदानुक्रम अनुसार पदनाम लिखें.
         """
        
    instruction += f"""
    4. अनिवार्य नीचे की संरचना (हस्ताक्षर एवं प्रतिलिपि) - सभी हस्ताक्षर ब्लॉक (Signature Blocks) को बिना किसी कोष्ठक के दाईं ओर (Right Side) रखें:
       
       - प्राथमिक हस्ताक्षर खंड (यह पत्र के मुख्य भाग के तुरंत बाद दाईं ओर संरेखित होना चाहिए):
{primary_sig}
       
       - प्रतिलिपि (Endorsement) अनुभाग (यह बाईं ओर होना चाहिए):
          पृ. क्रमांक / {branch} / 2026-27                             राजगढ़, दिनांक :- ..................
          प्रतिलिपि सूचनार्थ एवं आवश्यक कार्यवाही हेतु प्रेषित:-
          1. कलेक्टर महोदय, जिला राजगढ़ (म.प्र.) की ओर सादर सूचनार्थ.
          2. [वरिष्ठ कार्यालय/संबंधित विभाग] की ओर सूचनार्थ.
          
       - अंतिम हस्ताक्षर खंड (यह प्रतिलिपि अनुभाग के बिल्कुल नीचे अंत में दाईं ओर संरेखित होना चाहिए):
{subordinate_sig}
          
    5. फुटर (Footer):
       - पत्र के बिल्कुल अंत में नीचे बाईं ओर फाइल-पथ का संदर्भ लिखें:
         "{f_path}"
    """

    routing_instr = f"""
    6. बुद्धिमान विभाग राउटर (Intelligent Department Router) नियम:
       - आपको आने वाले पत्र के इनपुट दस्तावेज़ सामग्री (Extracted Content), अतिरिक्त निर्देश (Core Context/Instructions) या अपलोड किए गए चित्र के विषय/संदर्भ का सूक्ष्मता से विश्लेषण करना है.
       - आपको "District Level Policy/Data Collection" (जिला स्तरीय नीति/डेटा संग्रह) और "Block Level Execution" (विकासखंड स्तरीय निष्पादन) के बीच स्पष्ट रूप से अंतर करना होगा.
       
       पदानुक्रम और नोडल विभाग निर्धारण नियम (Hierarchy and Nodal Department Rules):
       - पत्र के "विषय" (Subject) और "संदर्भ" (Reference) को सर्वोच्च प्राथमिकता देकर पहले जिला (District) स्तर पर नोडल विभाग/कार्यालय का स्वतः निर्धारण करें.
       - सामान्य जिला स्तरीय जानकारी या डेटा संकलन के पत्रों को "मुख्य कार्यपालन अधिकारी, जनपद पंचायत" (Block/Janpad level) को संबोधित या रूट न करें. इन्हें केवल तब ही जनपद पंचायत (Janpad CEO) स्तर पर भेजा जाएगा, जब पत्र का विषय विशेष रूप से और केवल किसी एकल ग्राम पंचायत (single Gram Panchayat) की अधोसंरचना या मनरेगा (MGNREGA) कार्यों तक ही सीमित हो.
       
       निम्नलिखित administrative विषय-टू-विभाग मैपिंग (Topic-to-Department Mapping) पदानुक्रम का कड़ाई से पालन करें:
"""
    for topic, department in ROUTING_MAP.items():
        routing_instr += f"         * यदि पत्र का विषय '{topic}' से संबंधित है -> '{department}' को संबोधित करें.\n"
        
    routing_instr += """
       - यदि विषय/संदर्भ उपरोक्त श्रेणियों में से किसी से भी सीधे मेल नहीं खाता है, तो पत्र के विषय/संदर्भ के आधार पर स्वयं तार्किक रूप से एक उपयुक्त शासकीय विभाग या अधिकारी का निर्धारण करें (उदा. 'तहसीलदार, राजगढ़' या अन्य संबंधित विभाग).
       - *महत्वपूर्ण*: पत्र के प्रारूप में 'प्रति,' (To Address) खंड के अंतर्गत, स्वचालित रूप से इस प्रकार से विभाग/अधिकारी को संबोधित करें:
         प्रति,
              [पहचाना गया विभाग/अधिकारी का नाम]
              जिला राजगढ़ (म.प्र.)
       - यह 'प्रति,' खंड शासकीय प्रारूप के शीर्ष भाग (जावक क्रमांक और दिनांक) के ठीक नीचे और 'विषय:-' खंड के ठीक ऊपर होना चाहिए.
       - कार्यालयीन टिप्पणी (Notesheet) को छोड़कर अन्य सभी पत्राचार प्रकारों में यह 'प्रति' खंड भरना अनिवार्य है.
    """
    instruction += routing_instr

    instruction += """
    7. आउटपुट नियंत्रण:
       - केवल तैयार किया गया पत्र/टिप्पणी/आदेश ही आउटपुट में प्रदान करें. कोई अन्य बाहरी बातचीत, परिचय, या "यह रहा आपका प्रारूप" जैसी बातें न लिखें.
       - जहां भी अज्ञात आंकड़े या नाम हों, वहां उचित वर्गाकार कोष्ठकों जैसे [क्रमांक], [दिनांक], [नाम], [आंकड़ा] का प्रयोग करें.

    8. डायनेमिक ड्राफ्टिंग लॉजिक (Dynamic Drafting Logic) एवं डेटा प्लेसमेंट और फॉर्मेटिंग लॉक:
       - Dynamic Content Lock: सिस्टम को हमेशा केवल और केवल यूजर द्वारा वर्तमान में अपलोड की गई फ़ाइल (PDF/Image) का ही वास्तविक विषय, संदर्भ, विभाग, और अंदर का डेटा पढ़ना है।
       - Example vs Reality: कोड में जो 'राजगढ़ गोल्ड स्टैंडर्ड' प्रारूप दिया गया है, वह केवल विज़ुअल स्टाइल (हेडर डिजाइन, 1.0 स्पेसिंग, 1-टैब पैराग्राफ इंडेंट, राइट-साइड सील/साइन ब्लॉक) की नकल करने के लिए है। उदाहरण का टेक्स्ट (सांख्यिकी अधिकारी, विकासखंडवार आँकड़े) कभी भी नए पत्रों में रिपीट नहीं होना चाहिए।
       - Strict Input Parsing: जब नया आवक पत्र अपलोड हो, तो इनपुट डेटा के आधार पर प्रति, विषय, संदर्भ, और नीचे की तालिका/विवरण को पूरी तरह नया और तार्किक रूप से तैयार करो।
       - डेटा प्लेसमेंट (Data Placement): यूजर द्वारा 'Data to be Filled' (प्रेषित की जाने वाली वास्तविक जानकारी/आँकड़े) बॉक्स/अनुभाग में प्रदान की गई राजगढ़ जिले की वास्तविक जानकारी/संख्या/आंकड़ों या प्रपत्र के विवरण को पत्र के मुख्य भाग ("उपरोक्त विषयांतर्गत निर्देशानुसार लेख है कि...") में अत्यंत तार्किक और व्यवस्थित रूप से शामिल किया जाना चाहिए।
       - फॉर्मेटिंग लॉक (Formatting Lock): पूरा पत्र हमारे तय 'राजगढ़ गोल्ड स्टैंडर्ड' के नियमों (1.0 लाइन स्पेसिंग, 1-टैब पैराग्राफ इंडेंट, साइज 15 बोल्ड हेडर) के साथ अनिवार्य रूप से 1 ही पेज में फिट होना चाहिए।
    """
    return instruction

def create_docx(draft_text):
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Mangal'
    font.size = Pt(12)
    font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    def apply_formatting(p, before=0, after=6, line=1.15):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    p_h1 = doc.add_paragraph()
    p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_formatting(p_h1, 0, 2, 1.0)
    run_h1 = p_h1.add_run("कार्यालय कलेक्टर (जनजातीय कार्य तथा अनुसूचित जाति कल्याण विभाग)")
    run_h1.bold = True
    run_h1.underline = True
    run_h1.font.size = Pt(15)

    p_h2 = doc.add_paragraph()
    p_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_formatting(p_h2, 0, 6, 1.0)
    run_h2 = p_h2.add_run("जिला राजगढ़ (ब्यावरा) म.प्र.")
    run_h2.bold = True
    run_h2.underline = True
    run_h2.font.size = Pt(15)

    p_contact = doc.add_paragraph()
    apply_formatting(p_contact, 6, 4, 1.0)
    p_contact.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    p_contact.add_run("E-Mail ID: dotw.rjg@mp.gov.in\tPhone No. 07372-255263")

    pPr = p_contact._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    p_meta = doc.add_paragraph()
    apply_formatting(p_meta, 4, 12, 1.15) 
    p_meta.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    selected_br = st.session_state.get("selected_branch", "योजना")
    p_meta.add_run(f"क्रमांक /       / {selected_br} / 2026-27\tराजगढ़, दिनांक....................")

    lines = draft_text.split('\n') if isinstance(draft_text, str) else []
    
    has_endorsement_no = False
    for line in lines:
        text_check = line.strip()
        if text_check.startswith(("पृ. क्रमांक", "पृ.क्रमांक", "पृ. क्र.", "पृ.क्र.")):
            has_endorsement_no = True
            break

    in_address = False
    for line in lines:
        text = line.strip()
        if not text:
            continue
            
        text = text.replace("**", "")
        text = text.replace("---", "")
        
        if "कार्यालय कलेक्टर" in text or "E-Mail ID:" in text or "क्रमांक /" in text or "Phone No." in text or "जिला राजगढ़ (ब्यावरा)" in text or text.strip("_") == "" or text.strip("-") == "" or text.strip("=") == "":
            continue

        if "प्रतिलिपि सूचनार्थ" in text and not has_endorsement_no:
            p_endo = doc.add_paragraph()
            apply_formatting(p_endo, 24, 12, 1.15)
            p_endo.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
            selected_br = st.session_state.get("selected_branch", "योजना")
            p_endo.add_run(f"पृ. क्रमांक /       / {selected_br} / 2026-27\tराजगढ़, दिनांक :- ....................")
            has_endorsement_no = True

        p = doc.add_paragraph()
        
        if text.startswith("प्रति,"):
            apply_formatting(p, 6, 2, 1.15)
            p.add_run(text)
            in_address = True
            continue
            
        if text.startswith(("विषय", "संदर्भ")):
            in_address = False
            apply_formatting(p, 12, 6, 1.15)
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.first_line_indent = Inches(-1.0)
            
            text_clean = text.replace("ः-", ":-").replace(":", ":-").replace(":-", ":-")
            parts = text_clean.split(':-', 1)
            if len(parts) == 2:
                p.add_run(parts[0] + ":- ").bold = True
                # यह फिल्टर विषय-संदर्भ के फालतू डैश और स्टार को साफ रखेगा
                detail_clean = parts[1].strip().lstrip('-').lstrip('*').strip()
                p.add_run("\t" + detail_clean)
            else:
                p.add_run(text).bold = True
            continue
        if in_address:
            p.paragraph_format.left_indent = Inches(1.0)
            apply_formatting(p, 0, 2, 1.15)
            p.add_run(text)
            continue

        if text.startswith(('1.', '2.', '3.', '4.', '-', '*')):
            p.paragraph_format.left_indent = Inches(0.5)
            apply_formatting(p, 0, 4, 1.15)
            p.add_run(text)
            continue

        if text.startswith(("पृ. क्रमांक", "पृ.क्रमांक", "पृ. क्र.", "पृ.क्र.")) or "राजगढ़, दिनांक :-" in text:
            import re
            if "राजगढ़, दिनांक" in text:
                parts = re.split(r'\s*राजगढ़,\s*दिनांक', text)
                left_part = parts[0].strip()
                right_part = "राजगढ़, दिनांक" + parts[1].strip() if len(parts) > 1 else ""
                left_part = re.sub(r'\s+', ' ', left_part)
                text = f"{left_part}\t{right_part}"
            
            apply_formatting(p, 24, 12, 1.15)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
            p.add_run(text)
            continue
            
        if "प्रतिलिपि सूचनार्थ" in text:
            apply_formatting(p, 6, 6, 1.15)
            p.add_run(text).bold = True
            continue

        sig_keywords = [
            "कलेक्टर", "अपर कलेक्टर", "जिला संयोजक", 
            "अनुसूचित जाति तथा जनजातीय कल्याण विभाग", 
            "जिला राजगढ़ (म०प्र०)", "जिला राजगढ़ (म.प्र.)", "जिला राजगढ़"
        ]
        cleaned_text = text.strip("()[]{}* \t,.-")
        is_signature_line = False
        if text and "कार्यालय" not in text and "महोदय" not in text and "सूचनार्थ" not in text and "प्रतिलिपि" not in text and "निर्देश" not in text:
            if cleaned_text in sig_keywords or text in sig_keywords:
                is_signature_line = True

        if is_signature_line:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0)
            apply_formatting(p, 0, 0, 1.1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(4.75), docx.enum.text.WD_TAB_ALIGNMENT.CENTER)
            p.add_run("\t" + text).bold = True
            continue

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_formatting(p, 0, 8, 1.15)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.add_run(text)

    import io
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def render_draft_to_html(draft_text):
    """Renders the draft text into a styled HTML preview with right-aligned signature blocks."""
    sig_keywords = [
        "कलेक्टर", "अपर कलेक्टर", "जिला संयोजक", 
        "अनुसूचित जाति तथा जनजातीय कल्याण विभाग", 
        "जिला राजगढ़ (म०प्र०)", "जिला राजगढ़ (म.प्र.)", "जिला राजगढ़"
    ]
    
    lines = draft_text.split('\n')
    html_lines = []
    in_sig_block = False
    sig_block_lines = []
    
    def flush_sig_block(lines_list):
        if not lines_list:
            return ""
        joined_text = "<br>".join(lines_list)
        return f"""
        <div style="text-align: right; margin-top: 10px; margin-bottom: 10px;">
            <div style="display: inline-block; text-align: center; padding-right: 4rem; font-weight: bold; color: #7A1C1C; line-height: 1.4;">
                {joined_text}
            </div>
        </div>
        """

    has_endorsement_no = False
    for line in lines:
        trimmed_check = line.strip()
        if trimmed_check.startswith(("पृ. क्रमांक", "पृ.क्रमांक", "पृ. क्र.", "पृ.क्र.")):
            has_endorsement_no = True
            break

    in_address = False
    for line in lines:
        trimmed = line.strip()
        cleaned_trimmed = trimmed.strip("()[]{}* \t,.-")
        
        if trimmed and (trimmed.strip("_") == "" or trimmed.strip("-") == "" or trimmed.strip("=") == ""):
            continue
            
        is_sig = False
        if trimmed and "कार्यालय" not in trimmed and "महोदय" not in trimmed and "सूचनार्थ" not in trimmed and "प्रतिलिपि" not in trimmed and "निर्देश" not in trimmed:
            if cleaned_trimmed in sig_keywords or trimmed in sig_keywords:
                is_sig = True
            else:
                for kw in sig_keywords:
                    if cleaned_trimmed == kw.strip("()[]{}* \t,.-"):
                        is_sig = True
                        break
                        
        if is_sig:
            in_sig_block = True
            sig_block_lines.append(trimmed.replace("**", "").strip())
        else:
            if in_sig_block:
                html_lines.append(flush_sig_block(sig_block_lines))
                sig_block_lines = []
                in_sig_block = False
            
            if not trimmed:
                html_lines.append("<br>")
            elif "कार्यालय कलेक्टर" in trimmed or "जनजातीय कार्य" in trimmed:
                html_lines.append(f"<div style='text-align: center; font-weight: bold; font-size: 1.15rem; color: #7A1C1C; text-decoration: underline;'>{trimmed}</div>")
            elif "जिला राजगढ़" in trimmed and not is_sig:
                html_lines.append(f"<div style='text-align: center; font-weight: bold; font-size: 1.15rem; color: #7A1C1C;'>{trimmed}</div>")
            elif "E-Mail ID:" in trimmed and "Phone No." in trimmed:
                parts = trimmed.split("Phone No.")
                email_part = parts[0].strip()
                phone_part = "Phone No. " + parts[1].strip() if len(parts) > 1 else ""
                html_lines.append(f"""
                <div style="display: flex; justify-content: space-between; font-weight: bold; color: #7A1C1C; margin-top: 5px; margin-bottom: 5px;">
                    <div>{email_part}</div>
                    <div>{phone_part}</div>
                </div>
                <hr style="border: 1px solid #7A1C1C; margin-top: 5px; margin-bottom: 10px;">
                """)
            elif "E-Mail ID:" in trimmed or "__________" in trimmed or "|| कार्यालयीन टिप्पणी ||" in trimmed or "// कार्यालय आदेश //" in trimmed:
                html_lines.append(f"<div style='text-align: center; font-weight: bold; color: #7A1C1C;'>{trimmed}</div>")
            elif ("क्रमांक" in trimmed or "पृ. क्रमांक" in trimmed or "पृ.क्रमांक" in trimmed) and ("दिनांक" in trimmed or "राजगढ़, दिनांक" in trimmed):
                import re
                parts = re.split(r'\s*राजगढ़,\s*दिनांक', trimmed)
                left_part = parts[0].strip()
                right_part = "राजगढ़, दिनांक" + parts[1].strip() if len(parts) > 1 else ""
                html_lines.append(f"""
                <div style="display: flex; justify-content: space-between; margin-top: 5px; margin-bottom: 5px;">
                    <div>{left_part}</div>
                    <div>{right_part}</div>
                </div>
                """)
            elif "राजगढ़, दिनांक" in trimmed or "दिनांक:" in trimmed:
                html_lines.append(f"<div style='text-align: right; padding-right: 1rem;'>{trimmed}</div>")
            elif trimmed.startswith("प्रति,"):
                html_lines.append(f"<div style='font-weight: bold; color: #7A1C1C; margin-top: 8px;'>{trimmed}</div>")
                in_address = True
                continue
            elif trimmed.startswith(("विषय:-", "विषयः-", "संदर्भ:-", "संदर्भः-", "विषय:", "संदर्भ:")) or "विषय" in trimmed[:10] or "संदर्भ" in trimmed[:10]:
                in_address = False
                text_clean = trimmed.replace("ः-", ":-").replace(":", ":-").replace(":-", ":-")
                parts = text_clean.split(':-', 1)
                if len(parts) == 2:
                    html_lines.append(f"""
                    <div style="display: flex; margin-top: 6px; margin-bottom: 4px;">
                        <div style="font-weight: bold; color: #7A1C1C; min-width: 1.0in; max-width: 1.0in;">{parts[0]}:-</div>
                        <div style="text-align: left; flex-grow: 1;">{parts[1].strip()}</div>
                    </div>
                    """)
                else:
                    html_lines.append(f"<div style='font-weight: bold; color: #7A1C1C; margin-top: 8px;'>{trimmed}</div>")
                continue
            elif "प्रतिलिपि सूचनार्थ" in trimmed:
                in_address = False
                if not has_endorsement_no:
                    selected_br = st.session_state.get("selected_branch", "योजना")
                    html_lines.append(f"""
                    <div style="display: flex; justify-content: space-between; margin-top: 15px; margin-bottom: 10px; font-weight: normal;">
                        <div>पृ. क्रमांक /       / {selected_br} / 2026-27</div>
                        <div>राजगढ़, दिनांक :- ....................</div>
                    </div>
                    """)
                    has_endorsement_no = True
                html_lines.append(f"<div style='font-weight: bold; color: #7A1C1C; margin-top: 12px;'>{trimmed}</div>")
                continue
            elif in_address:
                html_lines.append(f"<div style='text-align: left; margin-left: 1.0in; margin-top: 2px; margin-bottom: 2px;'>{trimmed}</div>")
                continue
            else:
                formatted_line = trimmed
                parts = trimmed.split('**')
                if len(parts) > 1:
                    new_parts = []
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            new_parts.append(f"<b>{part}</b>")
                        else:
                            new_parts.append(part)
                    formatted_line = "".join(new_parts)
                html_lines.append(f"<div style='text-align: left; text-indent: 1.5rem;'>{formatted_line}</div>")
                
    if in_sig_block:
        html_lines.append(flush_sig_block(sig_block_lines))
        
    container_style = """
    <div style="background-color: white; border: 1px solid #E0DCD3; border-radius: 8px; padding: 30px; box-shadow: inset 0 0 10px rgba(0,0,0,0.02); font-family: 'Noto Serif Devanagari', serif; min-height: 400px; color: #333; line-height: 1.6;">
        {content}
    </div>
    """
    return container_style.format(content="".join(html_lines))

# ---------------------------------------------------------
# App Interface Layout (शासकीय प्रारूपक टैब्स - नो एरर लॉक)
# ---------------------------------------------------------
tab1, tab2 = st.tabs(["📝 शासकीय प्रारूपक (Draft Engine)", "🤖 जार्विस असिस्टेंट (Jarvis Mode)"])
col1, col2 = tab1.columns([1, 3])
col1, col2 = st.columns([1, 3])

with col1:
    st.markdown('<div class="css-card">', unsafe_allow_html=True)
    st.subheader("📥 इनपुट स्रोत (Input Source)")
    
    uploaded_file = st.file_uploader(
        "दस्तावेज़/चित्र अपलोड करें (Upload Document/Image)",
        type=["pdf", "txt", "png", "jpg", "jpeg"],
        help="संदर्भ दस्तावेज़ अपलोड करें जिसके आधार पर शासकीय आलेख तैयार किया जाना है।"
    )
    
    extracted_content = ""
    image_obj = None
    
    if uploaded_file is not None:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        
        if file_ext == "pdf":
            st.info("📄 पीडीएफ फाइल संसाधित की जा रही है...")
            extracted_content = extract_text_from_pdf(uploaded_file)
            if extracted_content:
                st.success("सफलतापूर्वक पीडीएफ से पाठ निकाल लिया गया है!")
                with st.expander("निकाला गया पाठ देखें (Extracted Text)"):
                    st.text_area("Extracted Text", extracted_content, height=150, disabled=True)
            
        elif file_ext == "txt":
            st.info("📝 टेक्स्ट फाइल संसाधित की जा रही है...")
            try:
                extracted_content = uploaded_file.read().decode("utf-8")
                st.success("सफलतापूर्वक टेक्स्ट पढ़ लिया गया है!")
                with st.expander("निकाला गया पाठ देखें (Extracted Text)"):
                    st.text_area("Extracted Text", extracted_content, height=150, disabled=True)
            except Exception as e:
                st.error(f"टेक्स्ट फ़ाइल पढ़ने में त्रुटि: {str(e)}")
                
        elif file_ext in ["png", "jpg", "jpeg"]:
            st.info("🖼️ चित्र संसाधित किया जा रहा है...")
            try:
                image_obj = Image.open(uploaded_file)
                st.image(image_obj, caption="अपलोड किया गया चित्र", use_column_width=True)
            except Exception as e:
                st.error(f"चित्र लोड करने में त्रुटि: {str(e)}")
                
    manual_context = st.text_area(
        "अतिरिक्त निर्देश / मुख्य विषय-वस्तु (Extra Instructions / Core Message)",
        height=150,
        placeholder="यहाँ वह मुख्य आदेश, विषय, अथवा निर्देश लिखें जिसके आधार पर शासकीय प्रारूप बनाना है..."
    )
    
    data_to_fill = st.text_area(
        "प्रेषित की जाने वाली वास्तविक जानकारी/आँकड़े (Data to be Filled)",
        height=150,
        placeholder="यहाँ राजगढ़ जिले की वास्तविक जानकारी, आंकड़े या प्रपत्र का विवरण दर्ज करें जो पत्र के मुख्य भाग में शामिल करना है..."
    )
    
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="css-card">', unsafe_allow_html=True)
    st.subheader("📤 प्रारूप निर्माण एवं संपादन (Draft Generation)")
    
    if "draft_text" not in st.session_state:
        st.session_state.draft_text = ""
        
    generate_btn = st.button("प्रारूप तैयार करें (Generate Draft)")
    
    if generate_btn:
        provider_val = st.session_state.get("model_provider_selection", "Google Gemini")
        has_key = True
        
        if provider_val == "Google Gemini":
            user_key = st.session_state.get("gemini_key_input", "").strip()
            api_key = user_key if user_key else os.environ.get("GEMINI_API_KEY")
            if not api_key:
                api_key = "AQ.Ab8RN6ImjKf-XDHlmdmt5LVTQmchwWaHWo-sygGk4pqDjB0kIg"
            
            if not api_key:
                st.error("❌ त्रुटि: कृपया एक मान्य गूगल जेमिनी API कुंजी प्रदान करें।")
                has_key = False
        elif provider_val == "Anthropic Claude":
            user_key = st.session_state.get("claude_key_input", "").strip()
            api_key = user_key if user_key else os.environ.get("ANTHROPIC_API_KEY")
            if not api_key:
                st.error("❌ त्रुटि: कृपया एक मान्य एंथ्रोपिक क्लॉड API कुंजी प्रदान करें।")
                has_key = False
        else:
            user_key = st.session_state.get("openai_key_input", "").strip()
            api_key = user_key if user_key else os.environ.get("OPENAI_API_KEY")
            if not api_key:
                st.error("❌ त्रुटि: कृपया एक मान्य OpenAI API कुंजी प्रदान करें।")
                has_key = False
        
        if has_key:
            with st.spinner("🤖 शासकीय प्रारूप तैयार किया जा रहा है, कृपया प्रतीक्षा करें..."):
                try:
                    system_inst = get_system_instruction(
                        workflow=workflow_type,
                        dept=dept_header,
                        district=district_info,
                        email=email_id,
                        branch=branch_name,
                        out_no=outward_no,
                        date=letter_date,
                        sign_off=signatory,
                        f_path=footer_path
                    )
                    
                    prompt = f"कृपया निम्न स्रोत सामग्री के आधार पर '{workflow_type}' का शासकीय प्रारूप तैयार करें। विशेष निर्देश:\n" \
                             f"1. पत्र के हेडर में जावक क्रमांक बिल्कुल हूबहू '{outward_no}' ही होना चाहिए, इसमें अपने मन से कोई शब्द न बदलें।\n" \
                             f"2. 'प्रति,' लिखने के बाद, उसके नीचे आने वाले प्राप्तकर्ता के पदनाम या विवरण को भी दो बार Tab स्पेस देकर उसी सीध में लाएं जहाँ से विषय का विवरण शुरू होता है।\n" \
                             f"3. 'विषय:-' और 'संदर्भ:-' लिखने के ठीक बाद दो बार Tab स्पेस (खाली जगह) दें, ताकि आगे का पूरा टेक्स्ट एक सीधी खड़ी लाइन (align) में व्यवस्थित दिखे।\n" \
                             f"4. प्राप्तकर्ता का पद, विषय का मैटर, and संदर्भ का मैटर—ये तीनों बिल्कुल एक ही वर्टिकल सीध (एक के नीचे एक) से शुरू होने चाहिए।\n" \
                             f"5. पत्र के मुख्य भाग (Body Text) की शुरुआत करते समय, संदर्भ का हवाला देने वाली पहली लाइन and विवरण वाली लाइन को आपस में जोड़कर एक ही निरंतर पैराग्राफ (continuous paragraph) में लिखें।\n" \
                             f"6. पत्र की भाषा या आवश्यक शासकीय सामग्री को जबरदस्ती छोटा या संकुचित (short) नहीं करना है। पत्र की सभी महत्वपूर्ण और जरूरी बातें पूरी गरिमा के साथ विस्तार से लिखी होनी चाहिए। बस लेआउट स्पेस का ऐसा उपयोग करें कि पूरा पत्र और प्रतिलिपि एक ही पेज (Single Page) पर आ सके।\n" \
                             f"7. *क्रिटिकल नियम (प्रति लॉक):* यदि पत्राचार प्रवाह (Workflow Type) 'अधीनस्थ को निर्देश/पत्र' है, तो 'बुद्धिमान विभाग राउटर' के नियमों को पूरी तरह बाईपास (ignore) करें। ऐसी स्थिति में पत्र को किसी जिला प्रमुख या वरिष्ठ कार्यालय को संबोधित करने के बजाय, सीधे अपलोड किए गए मूल दस्तावेज़ में निर्दिष्ट वास्तविक अधीनस्थ प्राप्तकर्ता के पदनाम और पते को ही 'प्रति,' में संबोधित करें।\n" \
                             f"8. *क्रिटिकल नियम (डेटा लॉक):* पत्र में दी गई वास्तविक तिथियां (Dates), समय-सीमा, ईमेल आईडी, पत्र क्रमांक और आंकड़ों को अपने मन से काल्पनिक रूप से न बदलें। अपलोड किए गए मूल पत्र में जो तथ्य, आधिकारिक ईमेल और तारीखें दी गई हैं, उन्हें पूरी शुद्धता के साथ हुबहू बनाए रखें।\n"
                    
                    if extracted_content:
                        prompt += f"--- स्रोत दस्तावेज़ सामग्री (Extracted Content) ---\n{extracted_content}\n\n"
                    
                    if manual_context:
                        prompt += f"--- मुख्य निर्देश (Core Context/Instructions) ---\n{manual_context}\n\n"
                        
                    if data_to_fill:
                        prompt += f"--- प्रेषित की जाने वाली वास्तविक जानकारी/आँकड़े (Data to be Filled) ---\n{data_to_fill}\n\n"
                        
                    if not extracted_content and not manual_context and not data_to_fill and not image_obj:
                        prompt += "नोट: कोई विशिष्ट स्रोत सामग्री नहीं दी गई है। कृपया मध्य प्रदेश शासन के पत्र प्रारूप का एक डमी उदाहरण तैयार करें।\n"
                        
                    model_name = selected_model
                    
                    if provider_val == "Google Gemini":
                        if HAS_NEW_SDK:
                            client = genai.Client(api_key=api_key)
                            contents = []
                            if image_obj is not None:
                                contents.append(image_obj)
                            contents.append(prompt)
                            
                            response = client.models.generate_content(
                                model=model_name,
                                contents=contents,
                                config=types.GenerateContentConfig(
                                    system_instruction=system_inst
                                )
                            )
                            st.session_state.draft_text = response.text
                        else:
                            genai_legacy.configure(api_key=api_key)
                            model = genai_legacy.GenerativeModel(
                                model_name=model_name,
                                system_instruction=system_inst
                            )
                            if image_obj is not None:
                                response = model.generate_content([prompt, image_obj])
                            else:
                                response = model.generate_content(prompt)
                            st.session_state.draft_text = response.text
                    elif provider_val == "Anthropic Claude":
                        client = anthropic.Anthropic(api_key=api_key)
                        messages_payload = []
                        
                        if image_obj is not None:
                            import base64
                            buffered = io.BytesIO()
                            img_format = image_obj.format if image_obj.format else "PNG"
                            mime_type = f"image/{img_format.lower()}"
                            if img_format.lower() == "jpg":
                                mime_type = "image/jpeg"
                            image_obj.save(buffered, format=img_format)
                            img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
                            
                            messages_payload.append({
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image",
                                        "source": {
                                            "type": "base64",
                                            "media_type": mime_type,
                                            "data": img_str,
                                        },
                                    },
                                    {
                                        "type": "text",
                                        "text": prompt
                                    }
                                ]
                            })
                        else:
                            messages_payload.append({
                                "role": "user",
                                "content": prompt
                            })
                            
                        response = client.messages.create(
                            model=model_name,
                            max_tokens=4000,
                            system=system_inst,
                            messages=messages_payload
                        )
                        st.session_state.draft_text = response.content[0].text
                    else:
                        client = OpenAI(api_key=api_key)
                        messages_payload = [
                            {"role": "system", "content": system_inst}
                        ]
                        
                        if image_obj is not None:
                            import base64
                            buffered = io.BytesIO()
                            img_format = image_obj.format if image_obj.format else "PNG"
                            mime_type = f"image/{img_format.lower()}"
                            if img_format.lower() == "jpg":
                                mime_type = "image/jpeg"
                            image_obj.save(buffered, format=img_format)
                            img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
                            
                            messages_payload.append({
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": prompt
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{mime_type};base64,{img_str}"
                                        }
                                    }
                                ]
                            })
                        else:
                            messages_payload.append({
                                "role": "user",
                                "content": prompt
                            })
                            
                        response = client.chat.completions.create(
                            model=model_name,
                            messages=messages_payload
                        )
                        st.session_state.draft_text = response.choices[0].message.content
                        
                    st.success("✔️ प्रारूप सफलतापूर्वक तैयार हो गया है!")
                    
                except Exception as e:
                    st.error(f"प्रारूप तैयार करने में त्रुटि: {str(e)}")
                    
    edited_draft = st.text_area(
        "संपादित करें (Edit Draft Text)",
        value=st.session_state.draft_text,
        height=400,
        help="आप इस प्रारूप को अपनी आवश्यकतानुसार संपादित कर सकते हैं।"
    )
    
    st.session_state.draft_text = edited_draft
    
    if st.session_state.draft_text:
        st.markdown("### 👀 प्रारूप पूर्वावलोकन (Draft Preview)")
        preview_html = render_draft_to_html(st.session_state.draft_text)
        st.markdown(preview_html, unsafe_allow_html=True)
    
    if st.session_state.draft_text:
        st.divider()
        st.subheader("📥 एक्सपोर्ट करें (Export)")
        
        doc_io = create_docx(st.session_state.draft_text)
        filename = f"{workflow_type.split(' ')[0]}{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        
        st.download_button(
            label="📄 एमएस वर्ड (.docx) फाइल डाउनलोड करें",
            data=doc_io,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    st.markdown('</div>', unsafe_allow_html=True)
# ---------------------------------------------------------
# Tab 2: जार्विस असिस्टेंट मोड (Iterative Conversational Chat)
# ---------------------------------------------------------
with tab2:
    st.markdown("<h3 style='color: #7A1C1C;'>🤖 जार्विस शासकीय सहायक (Jarvis Chat)</h3>", unsafe_allow_html=True)
    st.info("यहाँ आप जेमिनी से बोलकर या टाइप करके सीधे संवाद कर सकते हैं। यह आपके पिछले निर्देशों को याद रखेगा।")
    
    # चैट इतिहास को सुरक्षित रखने के लिए सेशन स्टेट
    if "jarvis_messages" not in st.session_state:
        st.session_state.jarvis_messages = []
        
    # पुराना चैट इतिहास स्क्रीन पर प्रदर्शित करना
    for msg in st.session_state.jarvis_messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            
    # 🎙️ जार्विस के लिए इनपुट विकल्प (वॉइस और टाइपिंग दोनों)
    jarvis_audio = st.audio_input("🎙️ जार्विस को बोलकर निर्देश दें (Speak to Jarvis)")
    jarvis_text = st.chat_input("या यहाँ टाइप करके जार्विस से बात करें...")
    
    user_speech_text = ""
    
    # यदि वॉइस इनपुट आया है तो उसे विस्पर से ट्रांसक्राइब करें
    if jarvis_audio is not None:
        with st.spinner("🔊 जार्विस आपकी आवाज़ सुन रहा है..."):
            try:
                user_openai_key = st.session_state.get("openai_key_input", "").strip()
                final_openai_key = user_openai_key if user_openai_key else os.environ.get("OPENAI_API_KEY")
                if final_openai_key:
                    client_ts = OpenAI(api_key=final_openai_key)
                    transcription = client_ts.audio.transcriptions.create(
                        model="whisper-1", 
                        file=jarvis_audio,
                        language="hi"
                    )
                    user_speech_text = transcription.text
                else:
                    st.warning("⚠️ जार्विस वॉइस टाइपिंग के लिए साइडबार में 'OpenAI API कुंजी' दर्ज करें।")
            except Exception as e:
                st.error(f"ट्रांसक्रिप्शन में त्रुटि: {str(e)}")

    # दोनों में से जो भी इनपुट मिला हो, उसे प्रोसेस करें
    final_input = jarvis_text if jarvis_text else user_speech_text
    
    if final_input:
        # यूजर का संदेश चैट में जोड़ें
        with st.chat_message("user"):
            st.write(final_input)
        st.session_state.jarvis_messages.append({"role": "user", "content": final_input})
        
        # जेमिनी चाबी का प्रबंध
        user_gemini_key = st.session_state.get("gemini_key_input", "").strip()
        final_gemini_key = user_gemini_key if user_gemini_key else os.environ.get("GEMINI_API_KEY")
        if not final_gemini_key:
            final_gemini_key = "AQ.Ab8RN6ImjKf-XDHlmdmt5LVTQmchwWaHWo-sygGk4pqDjB0kIg"
            
        with st.spinner("🤖 जार्विस नया प्रारूप तैयार कर रहा है..."):
            try:
                system_inst = get_system_instruction(
                    workflow=workflow_type, dept=dept_header, district=district_info,
                    email=email_id, branch=branch_name, out_no=outward_no,
                    date=letter_date, sign_off=signatory, f_path=footer_path
                )
                
                # चैट हिस्ट्री को जेमिनी के अनुकूल प्रारूप में ढालना
                if HAS_NEW_SDK:
                    client = genai.Client(api_key=final_gemini_key)
                    history_payload = []
                    for m in st.session_state.jarvis_messages[:-1]:
                        history_payload.append(types.Content(
                            role="user" if m["role"] == "user" else "model",
                            parts=[types.Part.from_text(text=m["content"])]
                        ))
                    
                    chat = client.chats.create(
                        model=selected_model,
                        config=types.GenerateContentConfig(system_instruction=system_inst, history=history_payload)
                    )
                    response = chat.send_message(final_input)
                    ai_response = response.text
                else:
                    genai_legacy.configure(api_key=final_gemini_key)
                    model = genai_legacy.GenerativeModel(model_name=selected_model, system_instruction=system_inst)
                    chat = model.start_chat(history=[])
                    response = chat.send_message(final_input)
                    ai_response = response.text
                
                # एआई का जवाब चैट में दिखाना और सहेजना
                with st.chat_message("assistant"):
                    st.write(ai_response)
                st.session_state.jarvis_messages.append({"role": "assistant", "content": ai_response})
                st.rerun()
                
            except Exception as e:
                st.error(f"जार्विस रिस्पॉन्स एरर: {str(e)}")
